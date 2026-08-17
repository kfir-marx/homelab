from __future__ import annotations

import hashlib
import os
from io import BytesIO
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Flowable, ListFlowable, Paragraph, SimpleDocTemplate, Spacer

from .interfaces import GenerationResult, StoredArtifact
from .security import UnsafeInput


class FilesystemArtifactStorage:
    def __init__(self, root: Path) -> None:
        self.root = root.resolve()

    def _path(self, key: str) -> Path:
        candidate = (self.root / key).resolve()
        if self.root not in candidate.parents:
            raise UnsafeInput("artifact key escapes storage root")
        return candidate

    def put(self, key: str, content: bytes, mime_type: str) -> StoredArtifact:
        target = self._path(key)
        target.parent.mkdir(parents=True, exist_ok=True, mode=0o700)
        temporary = target.with_suffix(target.suffix + ".tmp")
        with temporary.open("wb") as handle:
            handle.write(content)
        os.chmod(temporary, 0o600)
        temporary.replace(target)
        return StoredArtifact(key, hashlib.sha256(content).hexdigest(), len(content), mime_type)

    def get(self, key: str) -> bytes:
        return self._path(key).read_bytes()


def render_markdown(result: GenerationResult) -> bytes:
    lines = [result.professional_summary, "", "## Skills"]
    for group, skills in result.skill_groups.items():
        lines.append(f"- **{group}:** {', '.join(skills)}")
    lines.extend(["", "## Experience"])
    lines.extend(f"- {bullet.text}" for bullet in result.experience_bullets)
    if result.unsupported_requirements:
        lines.extend(["", "## Human-review gaps"])
        lines.extend(f"- {gap}" for gap in result.unsupported_requirements)
    return ("\n".join(lines) + "\n").encode()


def render_docx(result: GenerationResult, template: Path | None = None) -> bytes:
    document = Document(str(template)) if template and template.is_file() else Document()
    document.add_heading("Professional Summary", level=1)
    document.add_paragraph(result.professional_summary)
    document.add_heading("Skills", level=1)
    for group, skills in result.skill_groups.items():
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{group}: ").bold = True
        paragraph.add_run(", ".join(skills))
    document.add_heading("Selected Experience", level=1)
    for bullet in result.experience_bullets:
        document.add_paragraph(bullet.text, style="List Bullet")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def render_pdf(result: GenerationResult) -> bytes:
    output = BytesIO()
    styles = getSampleStyleSheet()
    story: list[Flowable] = [
        Paragraph("Professional Summary", styles["Heading1"]),
        Paragraph(result.professional_summary, styles["BodyText"]),
    ]
    story.extend([Spacer(1, 8), Paragraph("Skills", styles["Heading1"])])
    for group, skills in result.skill_groups.items():
        story.append(Paragraph(f"<b>{group}:</b> {', '.join(skills)}", styles["BodyText"]))
    story.extend([Spacer(1, 8), Paragraph("Selected Experience", styles["Heading1"])])
    story.append(
        ListFlowable(
            [Paragraph(item.text, styles["BodyText"]) for item in result.experience_bullets]
        )
    )
    SimpleDocTemplate(output, pagesize=A4, title="Tailored CV").build(story)
    return output.getvalue()
